"""Generate DOCX and PDF from VITEEZY_PROJECT_HANDOVER_DOCUMENT.md"""
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
MD_FILE = ROOT / "VITEEZY_PROJECT_HANDOVER_DOCUMENT.md"
DOCX_FILE = ROOT / "VITEEZY_PROJECT_HANDOVER_DOCUMENT.docx"
PDF_FILE = ROOT / "VITEEZY_PROJECT_HANDOVER_DOCUMENT.pdf"


def parse_markdown_lines(text: str):
    lines = text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_lines)))
        elif line.strip().startswith("- "):
            items = []
            while i < len(lines) and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                items.append(re.sub(r"^[\-\*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ul", items))
            continue
        elif re.match(r"^\d+\.\s", line.strip()):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ol", items))
            continue
        elif line.strip():
            para = [line.strip()]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", "-", "*", "```")) and not re.match(r"^\d+\.\s", lines[i].strip()):
                para.append(lines[i].strip())
                i += 1
            blocks.append(("p", " ".join(para)))
            continue
        i += 1
    return blocks


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def build_docx(blocks):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, content in blocks:
        if kind == "h1":
            p = doc.add_heading(strip_md_inline(content), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "h2":
            doc.add_heading(strip_md_inline(content), level=1)
        elif kind == "h3":
            doc.add_heading(strip_md_inline(content), level=2)
        elif kind == "p":
            doc.add_paragraph(strip_md_inline(content))
        elif kind == "ul":
            for item in content:
                doc.add_paragraph(strip_md_inline(item), style="List Bullet")
        elif kind == "ol":
            for item in content:
                doc.add_paragraph(strip_md_inline(item), style="List Number")
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.25)
        elif kind == "table":
            rows = []
            for tl in content:
                if re.match(r"^\|\s*[-:]+", tl):
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    table.rows[r].cells[c].text = strip_md_inline(cell)
    doc.save(DOCX_FILE)
    print(f"Created: {DOCX_FILE}")


def _ascii_safe(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2022": "*",
        "\u2192": "->",
        "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("ascii", errors="replace").decode("ascii")


def build_pdf(blocks):
    from fpdf import FPDF

    class PDF(FPDF):
        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    def write_wrapped(text, size=11, style="", indent=0):
        pdf.set_font("Helvetica", style=style, size=size)
        pdf.set_x(pdf.l_margin + indent)
        pdf.multi_cell(0, 6, _ascii_safe(strip_md_inline(text)))

    for kind, content in blocks:
        if kind == "h1":
            pdf.ln(4)
            write_wrapped(content, size=18, style="B")
            pdf.ln(2)
        elif kind == "h2":
            pdf.ln(3)
            write_wrapped(content, size=14, style="B")
            pdf.ln(1)
        elif kind == "h3":
            pdf.ln(2)
            write_wrapped(content, size=12, style="B")
        elif kind == "p":
            write_wrapped(content)
            pdf.ln(1)
        elif kind == "ul":
            for item in content:
                write_wrapped(f"  - {item}", indent=4)
            pdf.ln(1)
        elif kind == "ol":
            for idx, item in enumerate(content, 1):
                write_wrapped(f"  {idx}. {item}", indent=4)
            pdf.ln(1)
        elif kind == "code":
            pdf.set_font("Courier", size=9)
            pdf.multi_cell(0, 5, _ascii_safe(content))
            pdf.set_font("Helvetica", size=11)
            pdf.ln(1)
        elif kind == "table":
            for tl in content:
                if re.match(r"^\|\s*[-:]+", tl):
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                write_wrapped(" | ".join(cells), size=9)
            pdf.ln(1)

    pdf.output(PDF_FILE)
    print(f"Created: {PDF_FILE}")


def main():
    if not MD_FILE.exists():
        print(f"Missing markdown file: {MD_FILE}")
        sys.exit(1)

    text = MD_FILE.read_text(encoding="utf-8")
    blocks = parse_markdown_lines(text)

    try:
        from docx import Document  # noqa: F401
        build_docx(blocks)
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        build_docx(blocks)

  # Prefer Word-based conversion on Windows for higher-fidelity PDF
    if sys.platform == "win32" and DOCX_FILE.exists():
        try:
            from docx2pdf import convert
            convert(str(DOCX_FILE), str(PDF_FILE))
            print(f"Created: {PDF_FILE}")
            return
        except Exception as exc:
            print(f"docx2pdf unavailable ({exc}); falling back to fpdf2.")

    try:
        from fpdf import FPDF  # noqa: F401
        build_pdf(blocks)
    except ImportError:
        print("Installing fpdf2...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "fpdf2", "-q"])
        build_pdf(blocks)


if __name__ == "__main__":
    main()
